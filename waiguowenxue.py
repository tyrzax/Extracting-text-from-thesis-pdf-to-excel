# 《外国文学评论》（2000-2019）
# 《当代外国文学》（2000-2019）
import os, sys
import pdftotext
import xlrd
import xlwt
from xlutils.copy import copy
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
import re
from tkinter import filedialog
# from xpinyin import Pinyin
import numpy as np

def danwei(institution_txt):
    institution_txt_list = []
    institution_txt = institution_txt.replace(',','，').replace('。','，')
    if len(institution_txt.split('，')) > 1:
        institution_txt_list = institution_txt.split('，')
    else:
        institution_txt_list.append(institution_txt)
    institution_list = []
    institution_remove_list = []
    for word in institution_txt_list:
        if '研究所' in word:
            institution_list.append(word.split('研究所')[0] + '研究所')
        elif '系' in word and ('学院' in word or '大学' in word):
            institution_list.append(word.split('系')[0] + '系')
        elif '学院' in word:
            institution_list.append(word.split('学院')[0] + '学院')
        elif '大学' in word:
            institution_list.append(word.split('大学')[0] + '大学')
        if ('学报' in word or '出版社' in word or '科研' in word or '成果' in word) and word in institution_list:
            institution_list.remove(word)
    institution_list = list(set(institution_list))
    for word in institution_list:
        for key in institution_list:
            if (word != key) and (word in key):
                print(institution_list)
                print(word)
                print(key)
                institution_remove_list.append(word)
    institution_remove_list = list(set(institution_remove_list))
    print(institution_remove_list)
    for word in institution_remove_list:
        institution_list.remove(word)
    if institution_list == []:
        institution = '找不到【作者单位】！'
    else:
        institution = '；'.join(institution_list)
    return institution


def longest_substring(str1,str2):
    # 首先创建一个长宽分别为len(str1),len(str2)的二维数组
    record = np.zeros(shape=(len(str1),len(str2)))
    # 获取第一个字符串的长度
    str1_length = len(str1)
    # 获取第二个字符串的长度
    str2_length = len(str2)
    # 最大长度
    maxLen = 0
    # 结束的索引
    maxEnd = 0
    for i in range(str1_length):
        for j in range(str2_length):
            # 判断两个字符串对应的索引是否相同
            if str1[i] == str2[j]:
                # 判断是否是第一行或者是第一列
                if i==0 or j==0:
                    # 如果是则对应索引置一
                    record[i][j] = 1
                else:
                    # 如果不是对应的索引则为其左上角对应的元素加一
                    record[i][j] = record[i-1][j-1]+1
            else:
                # 如果字符串对应的元素不相同则置零
                record[i][j] = 0
            # 判断记录数值是否大于最大长度
            if record[i][j] > maxLen:
                # 如果是则充值最大长度
                maxLen = record[i][j]
                # 将结束索引置为i
                maxEnd = i
    # print(record)
    # 返回结果为，起始的索引及最大的长度
    return maxEnd-maxLen+1,maxEnd+1

def cn_authortitle(filename,filetxt,author_title_txt):
    # # 第一页和第二页
    # firstpage = ''
    # if '-----next page-----' in filetxt:
    #     firstpage = 'next page tag'.join(filetxt.split('-----next page-----')[0:2]).replace('\n','（换行标记）')
    #     nextpage_txt = re.findall('next page tag.*?（换行标记）（换行标记）',firstpage)[0]
    #     firstpage = firstpage.replace(nextpage_txt,'').replace('（换行标记）','\n')
    # else:
    #     firstpage = filetxt

    if '摘..要' in author_title_txt: 
        at_txt = author_title_txt.split('摘..要')[0]
        # if 'DOI' in at_txt:
        #     remove_doi_txt = at_txt.replace('\n','huanhangbiaoji')
        #     remove_doi_list = re.findall('DOI.*?huanhangbiaoji',remove_doi_txt)
        #     for word in remove_doi_list:
        #         word = word.replace('huanhangbiaoji','\n')
        #         at_txt = at_txt.replace(word,'')
        author_txt = at_txt.strip('\n').strip().split('\n')[-1]
        if len(author_txt.replace(' ','')) > 8:
            at_txt = at_txt.replace(author_txt,'')
            author_txt = at_txt.strip('\n').strip().split('\n')[-1]
        # author = author_txt
        # author = filename.split('_')[-1].replace('.txt','')
        title_txt = at_txt.replace(author_txt,'').replace('\n','').replace(' ','')
        title = title_txt
        # filename_title = filename.replace('_'+author+'.txt','')
        author = author_txt.replace(' ','')
        if ' ' in author_txt :
            author_list = author_txt.split(' ')
            author_list = [i for i in author_list if(len(str(i))!=0)]
            author = author_list[0]
            author2 = author_list[1]
        else:
            author2 = ''
        if len(author) < 2:
            author = author_txt.replace(' ','')
            author2 = ''
        # if author != author_txt:
        #     author2 = author_txt.replace(author,'')
        # else:
        #     author2 = ''

        # index
        # if '-' in filename_title:
        #     filename_title = filename_title.split('-')[1]
        # else:
        #     filename_title = filename_title.replace(filename_title.split('_')[0],'')
        # print(filename_title)
        # filename_index = filename_title[0:2]

        # if '_' in filename_index:
        #     if '_' == filename_index[0]:
        #         title_index = title_txt.rfind(filename_index[1])-1
        #         title = title_txt[title_index:]
        #     else:
        #         title_index = title_txt.rfind(filename_index[0])
        #         title = title_txt[title_index:]
        # else:
        #     title_index = title_txt.rfind(filename_index)
        #     title = title_txt[title_index:]
        # print(filename_index)
        # print(filename_title)

        # if len(title) < len(filename_title) and '_' == filename_title[0]:
        #     title = title_txt[-len(filename_title):] 
        #     if title[0] in '0123456789':
        #         for word in '0123456789':
        #             title = title.replace(word,'')
        # elif len(title) < len(filename_title) and len(title_txt.split(filename_index)) > 2:
        #     title = title_txt[-len(filename_title):] 


        # print(at_txt+'\nyes\n')
        # print(title_txt+'\nyes\n')
        # print(author_txt)

        # author_txt = at_txt.split(at_index)[1].replace('\n','').strip()
        # author = author_txt
        # title_txt = (at_txt.split(at_index)[0]+at_index).replace(' ','').replace('\n','')

        # print(title_txt)
        # print(author_txt)
        # if ' ' in author_txt :
        #     author_list = author_txt.split(' ')
        #     author_list = [i for i in author_list if(len(str(i))!=0)]
        #     author = author_list[0]
        #     author2 = author_list[1]
        # else:
        #     author2 = ''
        if len(author) < 2 or len(author) > 10:
            author = filename.split('_')[-1].replace('.txt','')
            author2 = ''

        # title_remove_list = ['理论研究', '二十世纪文学', '态势', '书评','学术争鸣', '动态', '讨论专栏', '经典作家研究史', '当代外国文学']
        # for word in title_remove_list:
        #     if word in title_txt[:7]:
        #         title_txt = title_txt.replace(word, '')
        # title = title_txt
    else:
        author = filename.split('_')[-1].replace('.txt','')
        author2 = ''
        if '-' in filename:
            title = filename.split('-')[-1].replace('_'+author,'').replace('.txt','')
        else:
            title = filename.replace('_'+author,'').replace('.txt','')

    if '”' in title and '“' not in title:
        title = '“' + title
    if '*' in title:
        title = title.replace('*', '')
    if '———' in title:
        title = title.replace('———', '——')
    if '——·' in title:
        title = title.replace('——·', '·')
    if '—' in title and '——' not in title:
        title = title.replace('—', '——')
    if '⊙' in title:
        title = title.replace('⊙', '')
    if len(title) < 2 or len(title) > 100:
        if '-' in filename:
            title = filename.split('-')[-1].replace('_'+author,'').replace('.txt','')
        else:
            title = filename.replace('_'+author,'').replace('.txt','')

    # print(title)
    if 'EN..TITLE' in filetxt and 'EN..ABSTRACT' in filetxt:
        en_title = filetxt.split('EN..TITLE')[1].split('EN..ABSTRACT')[0].replace('\n','')
    elif 'EN..TITLE' in filetxt:
        en_title = filetxt.split('EN..TITLE')[1].split('\n')[0]
    else:
        en_title = ''
    
    # author = filename.split('_')[-1].replace('.txt','')
    
    return author,author2,title,en_title


def country(file):
    with open(file, 'r') as f:
        txt = f.read()
    filename = file.split("/")[-1].split('.')[0]
    # print(filename)


    AsiaCountryList = ['中国', '日本', '韩国', '朝鲜', '蒙古', '越南', '老挝', '柬埔寨', '缅甸', '马来西亚', '新加坡', '泰国', '菲律宾', '印尼', '东帝汶', '巴基斯坦', '印度', '不丹', '尼泊尔', '孟加拉国', '马尔代夫', '斯里兰卡', '伊朗', '伊拉克', '科威特', '约旦', '沙特阿拉伯', '卡塔尔', '巴林', '阿联酋', '也门', '阿曼', '以色列', '阿塞拜疆', '格鲁吉亚', '亚美尼亚', '土耳其', '阿富汗', '黎巴嫩', '吉尔吉斯斯坦', '土库曼斯坦', '塔吉克斯坦', '乌兹别克斯坦', '文莱']
    WesternEuropeList = ['比利时', '法国', '爱尔兰', '卢森堡', '摩纳哥', '荷兰', '英国']
    NortheastEuropeList = ['白俄罗斯', '爱沙尼亚', '拉脱维亚', '立陶宛', '摩尔多瓦', '俄罗斯', '乌克兰', '丹麦', '芬兰', '冰岛', '瑞典']
    CentralandSouthernEuropeList = ['奥地利', '捷克', '德国', '匈牙利', '列支敦士登', '波兰', '斯洛伐克', '瑞士', '阿尔巴尼亚', '保加利亚', '克罗地亚', '希腊', '意大利', '马其顿', '马耳他', '葡萄牙', '罗马尼亚', '塞尔维亚', '斯洛文尼亚', '西班牙']
    AfricaList = ['阿尔及利亚', '安哥拉', '贝宁', '博茨瓦纳', '布基纳法索', '布隆迪', '佛得角', '喀麦隆', '中非', '乍得', '科摩罗', '科特迪瓦', '刚果民主共和国', '吉布提', '埃及', '赤道几内亚', '厄立特里亚', '埃塞俄比亚', '加蓬', '冈比亚', '加纳', '几内亚', '几内亚比绍', '肯尼亚', '莱索托', '利比里亚', '利比亚', '马达加斯加', '马拉维', '马里', '毛里塔尼亚', '毛里求斯', '马约特', '莫桑比克', '摩洛哥', '纳米比亚', '尼日尔', '尼日利亚', '刚果共和国', '卢旺达', '圣多美和普林西比', '塞内加尔', '塞舌尔', '塞拉利昂', '索马里', '南非', '南苏丹', '苏丹', '斯威士兰', '坦桑尼亚', '多哥', '突尼斯', '乌干达', '赞比亚', '津巴布韦']
    OceaniaList = ['澳大利亚', '巴布亚新几内亚', '新西兰', '斐济', '所罗门群岛', '瓦努阿图', '萨摩亚', '基里巴斯', '密克罗尼西亚联邦', '汤加', '马绍尔群岛', '帕劳', '图瓦卢', '瑙鲁']
    USList = ['美国']
    CanadaandotherAmericanCountriesList = ['加拿大', '墨西哥', '古巴', '海地', '多米尼加', '多米尼克', '格林纳达', '圣文森特和格林纳丁斯', '圣卢西亚', '圣基茨和尼维斯', '安提瓜和巴布达', '巴巴多斯', '伯利兹', '萨尔瓦多', '尼加拉瓜', '哥斯达黎加', '巴哈马', '巴拿马', '洪都拉斯', '哥伦比亚', '委内瑞拉', '玻利维亚', '秘鲁', '厄瓜多尔', '圭亚那', '苏里南', '巴西', '巴拉圭', '乌拉圭', '阿根廷', '智利', '危地马拉', '牙买加', '特立尼达和多巴哥']

    frequency_list = []
    # Asia
    AsiaCount = 0
    for AsiaCountry in AsiaCountryList:
        AsiaCount += len(re.findall(AsiaCountry,txt))
    # country_list.append('AsiaCountry:')
    # frequency_list.append(AsiaCount)
    frequency_list = {'亚洲':AsiaCount}

    # Western Europe
    WesternEuropeCount = 0
    for WesternEuropeCountry in WesternEuropeList:
        WesternEuropeCount += len(re.findall(WesternEuropeCountry,txt))
    # country_list.append('WesternEuropeCountry:')
    # frequency_list.append(WesternEuropeCount)
    frequency_list['西欧'] = WesternEuropeCount

    # Northeast Europe
    NortheastEuropeCount = 0
    for NortheastEuropeCountry in NortheastEuropeList:
        NortheastEuropeCount += len(re.findall(NortheastEuropeCountry,txt))
    # country_list.append('NortheastEuropeCountry:')
    # frequency_list.append(NortheastEuropeCount)
    frequency_list['东欧、北欧'] = NortheastEuropeCount

    # Central and Southern Europe
    CentralandSouthernEuropeCount = 0
    for CentralandSouthernEuropeCountry in CentralandSouthernEuropeList:
        CentralandSouthernEuropeCount += len(re.findall(CentralandSouthernEuropeCountry,txt))
    # country_list.append('CentralandSouthernEuropeCountry:')
    # frequency_list.append(CentralandSouthernEuropeCount)
    frequency_list['中欧、南欧'] = CentralandSouthernEuropeCount

    # Africa
    AfricaCount = 0
    for AfricaCountry in AfricaList:
        AfricaCount += len(re.findall(AfricaCountry,txt))
    # country_list.append('AfricaCountry:')
    # frequency_list.append(AfricaCount)
    frequency_list['非洲'] = AfricaCount

    # Oceania
    OceaniaCount = 0
    for OceaniaCountry in OceaniaList:
        OceaniaCount += len(re.findall(OceaniaCountry,txt))
    # country_list.append('OceaniaCountry:')
    # frequency_list.append(OceaniaCount)
    frequency_list['大洋洲'] = OceaniaCount

    # United States
    USCount = 0
    for USCountry in USList:
        USCount += len(re.findall(USCountry,txt))
    # country_list.append('USCountry:')
    # frequency_list.append(USCount)
    frequency_list['美国'] = USCount

    # Canada and other American Countries
    CanadaandotherAmericanCountriesCount = 0
    for CanadaandotherAmericanCountries in CanadaandotherAmericanCountriesList:
        CanadaandotherAmericanCountriesCount += len(re.findall(CanadaandotherAmericanCountries,txt))
    # country_list.append('CanadaandotherAmericanCountries:')
    # frequency_list.append(CanadaandotherAmericanCountriesCount)
    frequency_list['加拿大及其他美洲国家'] = CanadaandotherAmericanCountriesCount

    # count_list = ['亚洲:','西欧:','东欧、北欧:','中欧、南欧:','非洲:','大洋洲:','美国:','加拿大及其他美洲国家:',"未归类："]

    # print(frequency_list)
    max_frequency = max(frequency_list,key=frequency_list.get)
    # print(max_frequency)

    return max_frequency,frequency_list


def write_excel_xls_append(path, value):
    workbook = xlrd.open_workbook(path)
    sheets = workbook.sheet_names()
    worksheet = workbook.sheet_by_name(sheets[0])
    rows_old = worksheet.nrows
    new_workbook = copy(workbook)
    new_worksheet = new_workbook.get_sheet(0)
    new_worksheet.write(rows_old, 0, value)
    new_workbook.save(path)


def cutspace(ifn, ofn):
    infile = open(ifn, 'r')
    outfile = open(ofn, 'w')

    content = infile.read()
    content = content.replace('\n\n', '')
    content = content.replace('  ', '')
    outfile.write(content)

    infile.close
    outfile.close
    if (os.path.exists(ifn)):
        os.remove(ifn)
        # print '移除后test 目录下有文件：%s' %os.listdir(dirPath)
    else:
        print("要删除的文件不存在！")


def pdf2txt(pdfpath, txtpath):
    path = '/Users/YiFaye_Lee/Desktop/liuda/error_open.xls'
    with open(pdfpath, "rb") as f:
        try:
            pdf = pdftotext.PDF(f)
        except Exception:
            write_excel_xls_append(path, pdfpath + '无法打开')
            return

    with open(txtpath, 'w') as f:
        f.write("\n-----next page-----\n".join(pdf))

    # cutsapce
    # with open(txtpath.split('.')[0]+'-temp.txt', 'w') as f:
    #     f.write("\n-----next page-----\n".join(pdf))
    # cutspace(txtpath.split('.')[0]+'-temp.txt',txtpath)


def get_filelist(path, type):
    Filelist = []
    for home, dirs, files in os.walk(path):
        for filename in files:
            if filename.split('.')[-1] == type:
                Filelist.append(os.path.join(home, filename))
    return Filelist


def output(pdfpath):
    Filelist_pdf = get_filelist(pdfpath, 'pdf')
    Filelist_txt = get_filelist(pdfpath, 'txt')

    for filename in Filelist_pdf:
        if filename.replace('.pdf','.txt') in Filelist_txt:
            print('已存在!\n' + filename)
        else:
            print('正在处理:' + filename)
            print('还剩：' + str(len(Filelist_pdf) - len(get_filelist(pdfpath, 'txt')) - 1))
            pdf2txt(filename, filename.replace('.pdf','.txt'))
    print(len(Filelist_pdf))


def lunwen(filepath,m):
    filename = filepath.split("/")[-1]
    print(filename)
    with open(filepath, 'r') as f:
        filetxt = f.read()
    
    filetxt = filetxt.replace('＊', '*')
    zhaiyao_list = ['内 容提 要','摘      要 ：','摘    要 ：','[摘     要 ]','［摘        要］','［摘       要］','［摘      要］','［摘     要］','［摘    要］','［内容提要］', '［内容摘要］', '［提 要］', '［提要］', '［摘 要］', '［摘要］', '【内容提要】', '【内容摘要】', '【提 要】', '【提要】', '【摘 要】',
                    '摘       要：','【摘要】', '内容提要：', '内容提要:', '内容提要|', '内容摘要：', '内容摘要:', '内容摘要|', '提 要：', '提 要:', '提 要|', '提要：', '提要:',
                    '提    要：','提   要：','提  要：','提要|', '摘 要：', '摘 要:', '摘 要|', '摘要：', '摘要:', '摘要|', '内容提要', '内容摘要',
                    '〔内容提要〕', '〔内容摘要〕', '〔提 要〕', '〔提要〕', '〔摘 要〕', '〔摘要〕', '摘   要：', '摘   要:', '摘   要', '提   要 |',
                    '【提   要】', '提   要', '【提     要】', '提     要' ,'摘    要','摘       要','摘  要','摘     要','要］','摘      要：','摘     要：','摘    要：','摘  要：', '提 要', '提要', '摘 要', '摘要']
    guanjianci_list = ['关 键词','关 键 词','【关 键 词】','[ 关键词 ]','关键词：', '关键词:', '［关键词］', '关键词|', '【关键词】', '〔关键词〕', '关键词 |', '关键词 ：', '关键词 :', '关   键   词：','键词］','主题词','关键词']
    danwei_list = ['作 者 简 介：','[ 作者简介 ]','[作者单位]', '[作者简介]', '[作者信息]', '作者单位：', '作者单位:', '作者单位|', '作者单位 ：', '作者单位 :', '作者单位 |', '作者简介：',
                   '作者简介:', '作者简介|', '作者简介 ：', '作者简介 :', '作者简介 |', '作者信息：', '作者信息:', '作者信息|', '作者信息 ：', '作者信息 :',
                   '作者信息 |', '【作者单位】', '【作者简介】', '【作者信息】', '〔作者单位〕', '〔作者简介〕', '〔作者信息〕', '［作者简介］', '［作者单位］', '［作者信息］',
                   '作者简介：', '本 文 作 者 ：', '本文作者：','作者单位','作者简介']
    entitle_list = ['Title:','TITLE:','Title：','TITLE：'] 
    enzhaiyao_list = ['Abstract：','Abstract:','ABSTRACT：','ABSTRACT:']
    biaoti_list = ['标题：','标题:','题目：','题目:','标题 :','题目 :','题目 ：']
    enzuozhe_list = ['Authors: ','Author:']

    qikan_list = ['No.', 'NO.', 'DOI', '2018', '二○一八', '期', '辑', '第', '二零一八']
    # cut_sapce_list = filetxt.split('\n')
    # for word in cut_sapce_list:
    #     if '提' in word:
    #         filetxt = filetxt.replace(word,word.replace(' ',''))
    #     if '键' in word:
    #         filetxt = filetxt.replace(word,word.replace(' ',''))

    for word in zhaiyao_list:
        filetxt = filetxt.replace(word, '摘..要')
    for word in guanjianci_list:
        filetxt = filetxt.replace(word, "关..键..词")
    for word in danwei_list:
        filetxt = filetxt.replace(word, "单..位")
    for word in entitle_list:
        filetxt = filetxt.replace(word, "EN..TITLE")
    for word in biaoti_list:
        filetxt = filetxt.replace(word, "标..题")
    for word in enzuozhe_list:
        filetxt = filetxt.replace(word, "EN..AUTHOR")
    for word in enzhaiyao_list:
        filetxt = filetxt.replace(word,'EN..ABSTRACT')

    filetxt = filetxt.replace('(', '（')
    filetxt = filetxt.replace(')', '）')
    first_page = filetxt.split('-----next page-----')[0]

    # # 第一页和第二页
    page12 = filetxt
    if '-----next page-----' in filetxt:
        page12 = 'next page tag'.join(filetxt.split('-----next page-----')[0:2]).replace('\n','（换行标记）')
        nextpage_txt = re.findall('next page tag.*?（换行标记）（换行标记）',page12)[0]
        page12 = page12.replace(nextpage_txt,'').replace('（换行标记）','\n')

    author_title_txt = page12.replace(' ','')
    for word in zhaiyao_list:
        author_title_txt = author_title_txt.replace(word, '摘..要')
    author,author2,title,en_title = cn_authortitle(filename,filetxt,author_title_txt)

    # print(title)
    # print(en_title)
    # print(author)
    # print(author2)

    # 3. abstract
    abstract_txt = page12.replace(' ','')
    for word in zhaiyao_list:
        abstract_txt = abstract_txt.replace(word, '摘..要')
    if '摘..要' in abstract_txt and '关..键..词' in abstract_txt:
        abstract_txt = abstract_txt.split('摘..要')[1].split('关..键..词')[0]
        abstract = abstract_txt.replace('\n', '').replace(' ', '').replace('［', '').replace('］', '')
    elif '摘..要' in abstract_txt and '\n\n\n\n' in abstract_txt:
        abstract_txt = abstract_txt.split('摘..要')[1].split('\n\n\n\n')[0]
        abstract = abstract_txt.replace('\n', '').replace(' ', '').replace('［', '').replace('］', '')
    else:
        abstract = '找不到【摘要】！'
    if abstract[0] == ':' :
        abstract = abstract.replace(abstract[0],'')
            
    # print(abstract)

    # 4. keywords
    # keywords_txt = page12.replace(' ','')
    keywords_txt = page12
    if '关..键..词' in keywords_txt and 'DOI' in keywords_txt :
        keywords_txt = keywords_txt.split('关..键..词')[1].split('DOI')[0].replace('\n','')
        if '；' in keywords_txt:
            keywords = '；'.join(keywords_txt.replace('\n', '').replace(' ', '').split('；'))
        elif '，' in keywords_txt:
            keywords = '；'.join(keywords_txt.replace('\n', '').replace(' ', '').split('，'))
        else:
            keywords = '；'.join(keywords_txt.replace('\n', '').strip(' ').split())
        keywords = keywords.replace('，', '；')
    elif '关..键..词' in keywords_txt and '\n' in keywords_txt.split('关..键..词')[1] :
        keywords_txt = keywords_txt.split('关..键..词')[1].split('\n')[0].replace('\n','')
        if '；' in keywords_txt:
            keywords = '；'.join(keywords_txt.replace('\n', '').replace(' ', '').split('；'))
        elif '，' in keywords_txt:
            keywords = '；'.join(keywords_txt.replace('\n', '').replace(' ', '').split('，'))
        else:
            keywords = '；'.join(keywords_txt.replace('\n', '').strip(' ').split())
        keywords = keywords.replace('，', '；')
    else:
        keywords = '找不到【关键词】！'
    keywords.replace('；；','；')

    # 5. institution
    institution_txt = filetxt.replace(' ','')
    if '单..位' in institution_txt:
        institution_txt = filetxt.replace(' ','').split('单..位')[1].split('\n\n')[0].replace('\n', '')
    else:
        institution_txt = ''.join(institution_txt.split('-----next page-----')[-2:]).replace('\n','')
    if author2 != '':
        if '英文' in filename:
            if '。' in institution_txt:
                institution_txt1 = institution_txt.split('。')[0]
                institution_txt2 = institution_txt.split('。')[1]
            elif '；' in institution_txt:
                institution_txt1 = institution_txt.split('；')[0]
                institution_txt2 = institution_txt.split('；')[1]
            else:
                institution_txt1 = institution_txt
                institution_txt2 = ''
            institution = danwei(institution_txt1)
            institution2 = danwei(institution_txt2)
        else:
            if author2 in institution_txt:
                institution_txt1 = institution_txt.split(author2)[0]
                institution_txt2 = institution_txt.split(author2)[1]
            else:
                institution_txt1 = institution_txt
                institution_txt2 = ''
            institution = danwei(institution_txt1)
            institution2 = danwei(institution_txt2)
    else:
        institution =  danwei(institution_txt)
        institution2 = ''
    if institution2 !='' and institution2[0] == ',':
        institution2 = institution2.replace(institution2[0],'')
    


    # 6. journal
    journal_txt = filetxt.split('-----next page-----')[0:3]
    journal_txt = ''.join(journal_txt).replace(' ', '').split('\n')
    filename_journal = '当代外国文学'
    journal_txt_list = []
    journal_list = []
    for word in journal_txt:
        if filename_journal in word:
            journal_txt_list.append(word)
        if 'DOI' in word:
            journal_txt_list.append(word)
    # j name
    journal_list.append('《' + filename_journal + '》')
    # print(journal_txt)

    # j date
    # qikan_list = ['No.', 'No．', 'NO.', 'DOI', '2018', '二○一八', '期', '辑', '第', '二零一八']
    # for word in journal_txt_list:
    #     word = word.replace('，', ',').replace(' ', '').replace('．', '.').replace('·', '.')
    #     if '年第' in word:
    #         journal_date = re.findall("....年第.期", word)[0]
    #         journal_list.append(str(journal_date))
    #         break
    journal_year = filepath.split('/')[-2].split('-')[0]
    journal_date = filepath.split('/')[-2].split('-')[1]
    
    journal_list.append(journal_year + '年第' + journal_date + '期')

    journal = '，'.join(journal_list)
    # print(journal)

    # 7. project
    # project_txt = page12.replace(' ', '')
    project_txt = filetxt.replace(' ', '')
    xiangmu_list = ['*本文为', '*本文系', '*本文是', '本文系','本文为', '本文是', '本论文为', '本论文系', '本论文是','本文得到','本文已得到']
    # xiangmu_list = ['基金项目：','基金项目:','项目基金：','项目基金:']
    project = ''

    # 外国文学
    # for word in xiangmu_list:
    #     project_txt = project_txt.replace(word, "项..目")
    # if '项..目' in project_txt and '单..位' in project_txt:
    #     project = project_txt.split('项..目')[1].split('单..位')[0].replace(' ','').replace('\n','')
    # elif '项..目' in project_txt:
    #     project = project_txt.split('项..目')[1].split('\n')[0].replace(' ','')

    # 中外文学研究
    if re.findall("\*本文系.*?成果。", project_txt) != []:
        # print(re.findall("\*本文系.*?成果。", project_txt))
        project = project + re.findall("\*本文系.*?成果。", project_txt)[0]
    if re.findall('\*本文系.*?资助。', project_txt) != []:
        project = project + re.findall("\*本文系.*?资助。", project_txt)[0]

    if project != '' and project[0] == ':':
        project = project.replace(project[0],'')
    if project != '' and project[0] == '］':
        project = project.replace(project[0],'')
    for word in project:
        if word in '（）－：／０１２３４５６７８９ＡＢＣＤＥＦＧＨＩＪＫＬＭＮＯＰＲＳＴＵＶＷＸＹＺａｂｃｄｅｆｇｈｉｊｋｌｍｎｏｐｑｒｓｔｕｖｗｘｙｚ．，“”':
            project = project.replace(word,'')
    project = project.replace('*','').replace(' ','')

    # 8. country
    country_max,frequency_list = country(filepath)
    country_list = str(frequency_list)

    f.close()

    outputTable.write(m,0,title)
    outputTable.write(m,1,en_title)
    outputTable.write(m,2,author)
    outputTable.write(m,3,institution)
    outputTable.write(m,4,author2)
    outputTable.write(m,5,institution2)
    outputTable.write(m,6,keywords)
    outputTable.write(m,7,abstract)
    outputTable.write(m,8,project)
    outputTable.write(m,9,journal)
    outputTable.write(m,10,country_max)
    outputTable.write(m,11,country_list)



    # # docx
    # para_title = document.add_paragraph()
    # try:
    #     run_title = para_title.add_run(title)
    # except Exception:
    #     return
    # run_title.font.name = "黑体"
    # run_title._element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")

    # para_author = document.add_paragraph()
    # run_author = para_author.add_run('【作  者】' + author)
    # run_author.font.name = '宋体'
    # run_author._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")

    # para_institution = document.add_paragraph()
    # run_institution = para_institution.add_run('【单  位】' + institution)
    # run_institution.font.name = '宋体'
    # run_institution._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")

    # para_journal = document.add_paragraph()
    # run_journal = para_journal.add_run('【期  刊】' + journal)
    # run_journal.font.name = '宋体'
    # run_journal._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")

    # para_abstract = document.add_paragraph()
    # run_abstract = para_abstract.add_run('【内容摘要】' + abstract)
    # run_abstract.font.name = '宋体'
    # run_abstract._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")

    # para_keyword = document.add_paragraph()
    # run_keyword = para_keyword.add_run('【关键词】' + keywords)
    # run_keyword.font.name = '宋体'
    # run_keyword._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")

    # document.add_page_break()



# Find files that ends with .txt
def find(obj):
    if obj.endswith(".txt"):
        path_list.append(obj)


def get_list_dir(dir_path):
    dir_files = os.listdir(dir_path)
    for file in dir_files:
        file_path = os.path.join(dir_path, file)
        if os.path.isfile(file_path):
            find(file_path)
        if os.path.isdir(file_path):
            get_list_dir(file_path)




# Extract text first
pdfpath = filedialog.askdirectory(initialdir=sys.path[0],
                                         title="请选择你想要提取的文件夹")
output(pdfpath)

# Get those txt file dir
path_list=[]
get_list_dir(pdfpath)

# read error
readErrorfile = xlwt.Workbook(encoding = 'utf-8')
errtable = readErrorfile.add_sheet('data',cell_overwrite_ok=True)
errtable.write(0,0,"ioerror title")
errtable.write(0,1,"ioerror journal")
errtable.write(0,2,"ValueError title")
errtable.write(0,3,"ValueError journal")
errtable.write(0,4,"IndexError title")
errtable.write(0,5,"IndexError journal")
errtable.write(0,6,"会议信息及投稿指南")


# found error
# foundErrorfile = xlwt.Workbook(encoding = 'utf-8')
# errtable2 = foundErrorfile.add_sheet('data',cell_overwrite_ok=True)
# errtable.write(0,0,"论文名")
# errtable.write(0,1,"期刊名")
# errtable.write(0,2,"【作者】")
# errtable.write(0,3,"【标题】")
# errtable.write(0,4,"【关键词】")
# errtable.write(0,5,"【单位】")
# errtable.write(0,6,"【摘要】")

# Then find info
outputFile = xlwt.Workbook(encoding = 'utf-8')
outputTable = outputFile.add_sheet('data',cell_overwrite_ok=True)
outputTable.write(0,0,"中文题名")
outputTable.write(0,1,"英文题名")
outputTable.write(0,2,"第一作者（含职称职务等信息）")
outputTable.write(0,3,"机构")
outputTable.write(0,4,"第二作者（含职称职务等信息）")
outputTable.write(0,5,"机构")
outputTable.write(0,6,"关键词")
outputTable.write(0,7,"摘要")
outputTable.write(0,8,"项目")
outputTable.write(0,9,"期刊")
outputTable.write(0,10,"国别")
outputTable.write(0,11,"国别词频参考")


i = 1
j = 1
k = 1
m = 1
for path in path_list:
    try:
        filename = path.split("/")[-1]
        filename_journal = path.split('/')[-2]
        if '会议信息' in filename or '投稿指南' in filename:
            errtable.write(i,6,filename_journal+'/'+filename)
            i = i + 1
        else:
            lunwen(path,m)
            m = m + 1
            print("")
    except IOError as e:
        errtable.write(i,0,filename)
        errtable.write(i,1,filename_journal)
        i = i + 1
    # except ValueError as e:
    #     errtable.write(j,2,filename)
    #     errtable.write(j,3,filename_journal)
    #     j = j + 1
    except IndexError as e: 
        errtable.write(k,4,filename)
        errtable.write(k,5,filename_journal)
        k = k + 1

# document.save('/Users/YiFaye_Lee/Desktop/外国文学/nianjian2018/lunwen.docx')

outputFile.save('/Users/YiFaye_Lee/Desktop/liuda/output.xls')
readErrorfile.save('/Users/YiFaye_Lee/Desktop/liuda/error.xls')
# foundErrorfile.save('/Users/YiFaye_Lee/Desktop/外国文学/nianjian2018/error_found.xls')


