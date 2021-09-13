import os, sys
import pdftotext
import xlrd
import xlwt
from xlutils.copy import copy
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
import re
from tkinter import filedialog

def country(file):
    with open(file, 'r') as f:
        txt = f.read()
    filename = file.split("/")[-1].split('.')[0]
    # print(filename)


    AsiaCountryList = ['中国', '日本', '韩国', '朝鲜', '蒙古', '越南', '老挝', '柬埔寨', '缅甸', '马来西亚', '新加坡', '泰国', '菲律宾', '印尼', '东帝汶', '巴基斯坦', '印度', '不丹', '尼泊尔', '孟加拉国', '马尔代夫', '斯里兰卡', '伊朗', '伊拉克', '科威特', '约旦', '沙特阿拉伯', '卡塔尔', '巴林', '阿联酋', '也门', '阿曼', '以色列', '阿塞拜疆', '格鲁吉亚', '亚美尼亚', '土耳其', '阿富汗', '黎巴嫩', '吉尔吉斯斯坦', '土库曼斯坦', '塔吉克斯坦', '乌兹别克斯坦', '文莱']
    WesternEuropeList = ['比利时', '法国', '爱尔兰', '卢森堡', '摩纳哥', '荷兰', '英国']
    NortheastEuropeList = ['白俄罗斯', '爱沙尼亚', '拉脱维亚', '立陶宛', '摩尔多瓦', '俄罗斯', '乌克兰', '丹麦', '芬兰', '冰岛', '瑞典']
    CentralandSouthernEuropeList = ['奥地利', '捷克', '德国', '匈牙利', '列支敦士登', '波兰', '斯洛伐克', '瑞士', '阿尔巴尼亚', '保加利亚', '克罗地亚', '希腊', '意大利', '马其顿', '马耳他', '葡萄牙', '罗马尼亚', '塞尔维亚', '斯洛文尼亚', '西班牙']
    AfricaList = ['阿尔及利亚', '安哥拉', '贝宁', '博茨瓦纳', '布基纳法索', '布隆迪', '佛得角', '喀麦隆', '中非', '乍得', '科摩罗', '科特迪瓦', '刚果民主共和国', '吉布提', '埃及', '赤道几内亚', '厄立特里亚', '埃塞俄比亚', '加蓬', '冈比亚', '加纳', '几内亚', '几内亚比绍', '肯尼亚', '莱索托', '利比里亚', '利比亚', '马达加斯加', '马拉维', '马里', '毛里塔尼亚', '毛里求斯', '马约特', '莫桑比克', '摩洛哥', '纳米比亚', '尼日尔', '尼日利亚', '刚果共和国', '卢旺达', '圣多美和普林西比', '塞内加尔', '塞舌尔', '塞拉利昂', '索马里', '南非', '南苏丹', '苏丹', '斯威士兰', '坦桑尼亚', '多哥', '突尼斯', '乌干达', '赞比亚', '津巴布韦']
    OceaniaList = ['澳大利亚', '巴布亚新几内亚', '新西兰', '斐济', '所罗门群岛', '瓦努阿图', '萨摩亚', '基里巴斯', '密克罗尼西亚联邦', '汤加', '马绍尔群岛', '帕劳', '图瓦卢', '瑙鲁']
    USList = ['美国']
    CanadaandotherAmericanCountriesList = ['加拿大', '墨西哥', '古巴', '海地', '多米尼加', '多米尼克', '格林纳达', '圣文森特和格林纳丁斯', '圣卢西亚', '圣基茨和尼维斯', '安提瓜和巴布达', '巴巴多斯', '伯利兹', '萨尔瓦多', '尼加拉瓜', '哥斯达黎加', '巴哈马', '巴拿马', '洪都拉斯', '哥伦比亚', '委内瑞拉', '玻利维亚', '秘鲁', '厄瓜多尔', '圭亚那', '苏里南', '巴西', '巴拉圭', '乌拉圭', '阿根廷', '智利', '危地马拉', '牙买加', '特立尼达和多巴哥']

    frequency_list = []
    # Asia
    AsiaCount = 0
    for AsiaCountry in AsiaCountryList:
        AsiaCount += len(re.findall(AsiaCountry,txt))
    # country_list.append('AsiaCountry:')
    # frequency_list.append(AsiaCount)
    frequency_list = {'亚洲':AsiaCount}

    # Western Europe
    WesternEuropeCount = 0
    for WesternEuropeCountry in WesternEuropeList:
        WesternEuropeCount += len(re.findall(WesternEuropeCountry,txt))
    # country_list.append('WesternEuropeCountry:')
    # frequency_list.append(WesternEuropeCount)
    frequency_list['西欧'] = WesternEuropeCount

    # Northeast Europe
    NortheastEuropeCount = 0
    for NortheastEuropeCountry in NortheastEuropeList:
        NortheastEuropeCount += len(re.findall(NortheastEuropeCountry,txt))
    # country_list.append('NortheastEuropeCountry:')
    # frequency_list.append(NortheastEuropeCount)
    frequency_list['东欧、北欧'] = NortheastEuropeCount

    # Central and Southern Europe
    CentralandSouthernEuropeCount = 0
    for CentralandSouthernEuropeCountry in CentralandSouthernEuropeList:
        CentralandSouthernEuropeCount += len(re.findall(CentralandSouthernEuropeCountry,txt))
    # country_list.append('CentralandSouthernEuropeCountry:')
    # frequency_list.append(CentralandSouthernEuropeCount)
    frequency_list['中欧、南欧'] = CentralandSouthernEuropeCount

    # Africa
    AfricaCount = 0
    for AfricaCountry in AfricaList:
        AfricaCount += len(re.findall(AfricaCountry,txt))
    # country_list.append('AfricaCountry:')
    # frequency_list.append(AfricaCount)
    frequency_list['非洲'] = AfricaCount

    # Oceania
    OceaniaCount = 0
    for OceaniaCountry in OceaniaList:
        OceaniaCount += len(re.findall(OceaniaCountry,txt))
    # country_list.append('OceaniaCountry:')
    # frequency_list.append(OceaniaCount)
    frequency_list['大洋洲'] = OceaniaCount

    # United States
    USCount = 0
    for USCountry in USList:
        USCount += len(re.findall(USCountry,txt))
    # country_list.append('USCountry:')
    # frequency_list.append(USCount)
    frequency_list['美国'] = USCount

    # Canada and other American Countries
    CanadaandotherAmericanCountriesCount = 0
    for CanadaandotherAmericanCountries in CanadaandotherAmericanCountriesList:
        CanadaandotherAmericanCountriesCount += len(re.findall(CanadaandotherAmericanCountries,txt))
    # country_list.append('CanadaandotherAmericanCountries:')
    # frequency_list.append(CanadaandotherAmericanCountriesCount)
    frequency_list['加拿大及其他美洲国家'] = CanadaandotherAmericanCountriesCount

    # count_list = ['亚洲:','西欧:','东欧、北欧:','中欧、南欧:','非洲:','大洋洲:','美国:','加拿大及其他美洲国家:',"未归类："]

    # print(frequency_list)
    max_frequency = max(frequency_list,key=frequency_list.get)
    # print(max_frequency)

    return max_frequency,frequency_list


def write_excel_xls_append(path, value):
    workbook = xlrd.open_workbook(path)
    sheets = workbook.sheet_names()
    worksheet = workbook.sheet_by_name(sheets[0])
    rows_old = worksheet.nrows
    new_workbook = copy(workbook)
    new_worksheet = new_workbook.get_sheet(0)
    new_worksheet.write(rows_old, 0, value)
    new_workbook.save(path)


def cutspace(ifn, ofn):
    infile = open(ifn, 'r')
    outfile = open(ofn, 'w')

    content = infile.read()
    content = content.replace('\n\n', '')
    content = content.replace('  ', '')
    outfile.write(content)

    infile.close
    outfile.close
    if (os.path.exists(ifn)):
        os.remove(ifn)
        # print '移除后test 目录下有文件：%s' %os.listdir(dirPath)
    else:
        print("要删除的文件不存在！")


def pdf2txt(pdfpath, txtpath):
    path = '/Users/YiFaye_Lee/Desktop/外国文学/nianjian2018/error_open.xls'
    with open(pdfpath, "rb") as f:
        try:
            pdf = pdftotext.PDF(f)
        except Exception:
            write_excel_xls_append(path, pdfpath + '无法打开')
            return

    with open(txtpath, 'w') as f:
        f.write("\n-----next page-----\n".join(pdf))

    # cutsapce
    # with open(txtpath.split('.')[0]+'-temp.txt', 'w') as f:
    #     f.write("\n-----next page-----\n".join(pdf))
    # cutspace(txtpath.split('.')[0]+'-temp.txt',txtpath)


def get_filelist(path, type):
    Filelist = []
    for home, dirs, files in os.walk(path):
        for filename in files:
            if filename.split('.')[-1] == type:
                Filelist.append(os.path.join(home, filename))
    return Filelist


def output(pdfpath):
    Filelist_pdf = get_filelist(pdfpath, 'pdf')
    Filelist_txt = get_filelist(pdfpath, 'txt')

    for filename in Filelist_pdf:
        if filename.split('.')[0] + '.txt' in Filelist_txt:
            print('已存在!\n' + filename)
        else:
            print('正在处理:' + filename)
            print('还剩：' + str(len(Filelist_pdf) - len(get_filelist(pdfpath, 'txt')) - 1))
            pdf2txt(filename, filename.split('.')[0] + '.txt')
    print(len(Filelist_pdf))


def lunwen(filepath,m):
    filename = filepath.split("/")[-1]
    print(filename)
    with open(filepath, 'r') as f:
        filetxt = f.read()

    filetxt = filetxt.replace('＊', '*')
    zhaiyao_list = ['摘      要 ：','摘    要 ：','[摘     要 ]','［摘        要］','［摘       要］','［摘      要］','［摘     要］','［摘    要］','［内容提要］', '［内容摘要］', '［提 要］', '［提要］', '［摘 要］', '［摘要］', '【内容提要】', '【内容摘要】', '【提 要】', '【提要】', '【摘 要】',
                    '摘       要：','【摘要】', '内容提要：', '内容提要:', '内容提要|', '内容摘要：', '内容摘要:', '内容摘要|', '提 要：', '提 要:', '提 要|', '提要：', '提要:',
                    '提    要：','提   要：','提  要：','提要|', '摘 要：', '摘 要:', '摘 要|', '摘要：', '摘要:', '摘要|', '内容提要', '内容摘要', '提 要', '提要', '摘 要', '摘要',
                    '〔内容提要〕', '〔内容摘要〕', '〔提 要〕', '〔提要〕', '〔摘 要〕', '〔摘要〕', '摘   要：', '摘   要:', '摘   要', '提   要 |',
                    '【提   要】', '提   要', '【提     要】', '提     要' ,'摘    要','摘       要','摘  要','摘     要','要］','摘      要：','摘     要：','摘    要：','摘  要：']
    guanjianci_list = ['关 键 词','【关 键 词】','[ 关键词 ]','关键词：', '关键词:', '［关键词］', '关键词|', '【关键词】', '〔关键词〕', '关键词 |', '关键词 ：', '关键词 :', '关   键   词：','键词］','关键词']
    danwei_list = ['[ 作者简介 ]','[作者单位]', '[作者简介]', '[作者信息]', '作者单位：', '作者单位:', '作者单位|', '作者单位 ：', '作者单位 :', '作者单位 |', '作者简介：',
                   '作者简介:', '作者简介|', '作者简介 ：', '作者简介 :', '作者简介 |', '作者信息：', '作者信息:', '作者信息|', '作者信息 ：', '作者信息 :',
                   '作者信息 |', '【作者单位】', '【作者简介】', '【作者信息】', '〔作者单位〕', '〔作者简介〕', '〔作者信息〕', '［作者简介］', '［作者单位］', '［作者信息］',
                   '作者简介：', '本 文 作 者 ：', '本文作者：','作者单位','作者简介']
    entitle_list = ['Title:','TITLE:'] 
    # xiangmu_list = ['基金项目：','基金项目:','基金项目','*本文为','本文系','本文是','本论文为','本论文系','本论文是']
    # enzhaiyao_list = ['ABSTRACT:','Abstract:']
    qikan_list = ['No.', 'NO.', 'DOI', '2018', '二○一八', '期', '辑', '第', '二零一八']

    for word in zhaiyao_list:
        filetxt = filetxt.replace(word, '摘..要')
    for word in guanjianci_list:
        filetxt = filetxt.replace(word, "关..键..词")
    for word in danwei_list:
        filetxt = filetxt.replace(word, "单..位")
    for word in entitle_list:
        filetxt = filetxt.replace(word, "EN..TITLE")
    # for word in enzhaiyao_list:
    #     filetxt = filetxt.replace(word,'EN..ABSTRACT')
    filetxt = filetxt.replace('(', '（')
    filetxt = filetxt.replace(')', '）')
    first_page = filetxt.split('-----next page-----')[0]
    # print(first_page)

    # 1. author
    author = filename.split('-')[-1].split('.')[0]
    if '_' in author:
        author2 = author.split('_')[1]
        author = author.split('_')[0]
        # if len(author) = :
        #     pass
        # author = '，'.join(author)
    else:
        author2 = ''
    if len(author) > 10 or len(author) < 2:
        author = '找不到【作者】！'
        author2 = ''

    # print('【作者】' + author)

    # 2. title
    filename_title = filename.split('/')[-1].split('-')[0]
    title_txt = filetxt.split('-----next page-----')[0].replace(' ', '').replace('，', '')
    title_txt = title_txt.split(author.replace('，', ''))[0].replace(' ', '').replace('\n', '')
    # print(title_txt)
    if filename_title[0:1] == filename_title[1:2] and filename_title[1:2] == filename_title[2:3]:
        filename_title = filename_title[0:3]
    elif filename_title[0:1] == filename_title[1:2]:
        filename_title = filename_title[0:2]
    else:
        filename_title = filename_title[0:1]
        # print(filename_title_index)
    # print(filename_title)

    title_index = title_txt.rfind(filename_title)
    title = title_txt[title_index:]

    if '”' in title and '“' not in title:
        title = '“' + title
    if '*' in title:
        title = title.replace('*', '')
    if '———' in title:
        title = title.replace('———', '——')
    if '——·' in title:
        title = title.replace('——·', '·')
    if '—' in title and '——' not in title:
        title = title.replace('—', '——')
    if '⊙'  in title:
        title = title.replace('⊙','')
    if len(title)<2 or len(title) > 100:
        title = filename.split('/')[-1].split('-')[0]


    # print(title)

    if 'EN..TITLE' in filetxt:
        en_title = filetxt.split('EN..TITLE')[1].split('\n')[0]
    else:
        en_title = ''


    # 3. abstract
    abstract_txt = filetxt.replace(' ','')
    if '摘..要' in abstract_txt and '关..键..词' in abstract_txt:
        abstract_txt = abstract_txt.split('摘..要')[1].split('关..键..词')[0]
        abstract = abstract_txt.replace('\n', '').replace(' ', '').replace('［', '').replace('］', '')
    elif '摘..要' in abstract_txt and '\n\n\n\n' in abstract_txt:
        abstract_txt = abstract_txt.split('摘..要')[1].split('\n\n\n\n')[0]
        abstract = abstract_txt.replace('\n', '').replace(' ', '').replace('［', '').replace('］', '')
    else:
        abstract = '找不到【摘要】！'
            
    # print(abstract)

    # 4. keywords
    # if '关..键..词' in first_page and '中图分类号' in first_page.split('关..键..词')[1]:
    #     keywords_txt = first_page.split('关..键..词')[1].split('中图分类号')[0]
    #     keywords = '；'.join(keywords_txt.replace('\n','').strip(' ').split())
    # elif '关..键..词' in first_page and 'DOI' in first_page.split('关..键..词')[1]:
    #     keywords_txt = first_page.split('关..键..词')[1].split('DOI')[0]
    #     keywords = '；'.join(keywords_txt.replace('\n','').strip(' ').split())
    # elif '关..键..词' in first_page and '单..位' in first_page.split('关..键..词')[1]:
    #     keywords_txt = first_page.split('关..键..词')[1].split('单..位')[0]
    #     keywords = '；'.join(keywords_txt.replace('\n','').strip(' ').split())
    # elif '关..键..词' in first_page and '\n\n' in first_page.split('关..键..词')[1]:
    #     keywords_txt = first_page.split('关..键..词')[1].split('\n\n')[0]
    #     keywords = '；'.join(keywords_txt.replace('\n','').strip(' ').split())
    if '关..键..词' in first_page and '\n' in first_page.split('关..键..词')[1]:
        keywords_txt = first_page.split('关..键..词')[1].split('\n')[0]
        if '；' in keywords_txt:
            keywords = '；'.join(keywords_txt.replace('\n', '').replace('；', ' ').strip(' ').split())
        elif '，' in keywords_txt:
            keywords = '；'.join(keywords_txt.replace('\n', '').replace('，', ' ').strip(' ').split())
        else:
            keywords = '；'.join(keywords_txt.replace('\n', '').strip(' ').split())
        keywords = keywords.replace('，', '；')
    else:
        keywords = '找不到【关键词】！'
    if ']；' in keywords:
        keywords = keywords.replace(']；','')
    # if '基金项目' in keywords:

    # 5. institution
    # print( '单..位' in filetxt)
    institution_key = ['学院', '学校', '大学', '研究所']
    institution = ''
    # if author2 == '':
    if author[-2:] + '（' in first_page.replace('\n', '').replace(' ', ''):
        institution_txt = first_page.replace('\n', '').replace(' ', '').split(author[-2:] + '（')[1].split('）')[0]
        # print(institution_txt)
        institution_txt_list = []
        if len(institution_txt.split('，')) > 1:
            institution_txt_list = institution_txt.split('，')
        else:
            institution_txt_list.append(institution_txt)
        institution_list = []
        for word in institution_txt_list:
            word = word.replace(' ', '')
            if '研究所' in word:
                institution_list.append(word.split('研究所')[0] + '研究所')
            elif '系' in word and ('学院' in word or '大学' in word):
                institution_list.append(word.split('系')[0] + '系')
            elif '学院' in word:
                institution_list.append(word.split('学院')[0] + '学院')
            elif '大学' in word:
                institution_list.append(word.split('大学')[0] + '大学')
            if ('学报' in word or '出版社' in word or '科研' in word or '成果' in word) and word in institution_list:
                institution_list.remove(word)
        if institution_list == []:
            institution = '找不到【作者单位】！'
            institution2 = ''
        else:
            if author2 == '':
                institution = institution_list[0]
                institution2 = ''
            elif author2 != '' and len(institution_list)==2 :
                institution = institution_list[0]
                institution2 = institution_list[1]
            else:
                institution = '；'.join(institution_list)
                institution2 = ''
    elif '单..位' in filetxt:
        # institution_txt = filetxt.split('单..位')[1].split('\n\n')[0].replace('\n', '').replace(' ', '')
        if author2 == '':
            institution_txt = filetxt.split('单..位')[1].split('\n\n')[0].replace('\n', '').replace(' ', '')
            institution_txt_list = []
            # print(institution_txt)
            if len(institution_txt.split('，')) > 1:
                institution_txt_list = institution_txt.split('，')
            else:
                institution_txt_list.append(institution_txt)
            # institution_key = ['学院','学校','大学','研究所']
            institution_list = []
            for word in institution_txt_list:
                if '研究所' in word:
                    institution_list.append(word.split('研究所')[0] + '研究所')
                elif '系' in word and ('学院' in word or '大学' in word):
                    institution_list.append(word.split('系')[0] + '系')
                elif '学院' in word:
                    institution_list.append(word.split('学院')[0] + '学院')
                elif '大学' in word:
                    institution_list.append(word.split('大学')[0] + '大学')
                if ('学报' in word or '出版社' in word or '科研' in word or '成果' in word) and word in institution_list:
                    institution_list.remove(word)
            for word in institution_list:
                for key in institution_list:
                    if (word != key) and (word in key):
                        print(institution_list)
                        print(word)
                        print(key)
                        institution_list.remove(word)
            if institution_list == []:
                institution = '找不到【作者单位】！'
                institution2 = ''
            else:
                institution = '；'.join(institution_list)
                institution2 = ''
        else:
            institution_txt = filetxt.split('单..位')[1].split('\n\n')[0].replace('\n', '').replace(' ', '')
            institution_txt1 = institution_txt.split(author2)[0]
            if len(institution_txt.split(author2))>1:
                institution_txt2 = institution_txt.split(author2)[1]
            elif author2 in filetxt.split('单..位')[1]:
                institution_txt2 = filetxt.split('单..位')[1].replace(' ', '').split(author2)[1].split('\n\n')[0].replace('\n', '')
            else:
                institution_txt2 = ''
            institution_txt_list = []
            # print(institution_txt)
            if len(institution_txt1.split('，')) > 1:
                institution_txt_list = institution_txt1.split('，')
            else:
                institution_txt_list.append(institution_txt1)
            # institution_key = ['学院','学校','大学','研究所']
            institution_list = []
            for word in institution_txt_list:
                if '研究所' in word:
                    institution_list.append(word.split('研究所')[0] + '研究所')
                elif '系' in word and ('学院' in word or '大学' in word):
                    institution_list.append(word.split('系')[0] + '系')
                elif '学院' in word:
                    institution_list.append(word.split('学院')[0] + '学院')
                elif '大学' in word:
                    institution_list.append(word.split('大学')[0] + '大学')
                if ('学报' in word or '出版社' in word or '科研' in word or '成果' in word) and word in institution_list:
                    institution_list.remove(word)
            for word in institution_list:
                for key in institution_list:
                    if word != key and word in key:
                        institution_list.remove(word)
            if institution_list == []:
                institution = '找不到【作者单位】！'
            else:
                institution = '；'.join(institution_list)

            institution_txt_list = []
            # print(institution_txt)
            if len(institution_txt2.split('，')) > 1:
                institution_txt_list = institution_txt2.split('，')
            else:
                institution_txt_list.append(institution_txt2)
            # institution_key = ['学院','学校','大学','研究所']
            institution_list = []
            for word in institution_txt_list:
                if '研究所' in word:
                    institution_list.append(word.split('研究所')[0] + '研究所')
                elif '系' in word and ('学院' in word or '大学' in word):
                    institution_list.append(word.split('系')[0] + '系')
                elif '学院' in word:
                    institution_list.append(word.split('学院')[0] + '学院')
                elif '大学' in word:
                    institution_list.append(word.split('大学')[0] + '大学')
                if ('学报' in word or '出版社' in word or '科研' in word or '成果' in word) and word in institution_list:
                    institution_list.remove(word)
            for word in institution_list:
                for key in institution_list:
                    if word != key and word in key:
                        institution_list.remove(word)
            if institution_list == []:
                institution2= '找不到【作者单位】！'
            else:
                institution2 = '；'.join(institution_list)
    else:
        institution_txt_list = []
        institution_list = []
        for word in first_page.split('\n'):
            if '，' in word:
                word = word.split('，')
                for key in word:
                    institution_txt_list.append(key)
            else:
                institution_txt_list.append(word)
        for word in institution_txt_list:
            word = word.replace(' ', '')
            if '研究所' in word:
                institution_list.append(word.split('研究所')[0] + '研究所')
            elif '系' in word and ('学院' in word or '大学' in word):
                institution_list.append(word.split('系')[0] + '系')
            elif '学院' in word:
                institution_list.append(word.split('学院')[0] + '学院')
            elif '大学' in word:
                institution_list.append(word.split('大学')[0] + '大学')
                institution_list.append(word.split('出版社')[0] + '出版社')
            if ('学报' in word or '出版社' in word or '科研' in word or '成果' in word) and word in institution_list:
                institution_list.remove(word)
        if institution_list == []:
            institution = '找不到【作者单位】！'
            institution2 = ''
        else:
            institution = '；'.join(institution_list)
            institution2 = ''
    if institution == '找不到【作者单位】！':
        institution_txt_list = []
        institution_list = []
        for word in first_page.split('\n'):
            if '，' in word:
                word = word.split('，')
                for key in word:
                    institution_txt_list.append(key)
            else:
                institution_txt_list.append(word)
        for word in institution_txt_list:
            word = word.replace(' ', '')
            if '研究所' in word:
                institution_list.append(word.split('研究所')[0] + '研究所')
            elif '系' in word and ('学院' in word or '大学' in word):
                institution_list.append(word.split('系')[0] + '系')
            elif '学院' in word:
                institution_list.append(word.split('学院')[0] + '学院')
            elif '大学' in word:
                institution_list.append(word.split('大学')[0] + '大学')
                institution_list.append(word.split('出版社')[0] + '出版社')
            if ('学报' in word or '出版社' in word or '科研' in word or '成果' in word) and word in institution_list:
                institution_list.remove(word)
        if institution_list == []:
            institution = '找不到【作者单位】！'
            institution2 = ''
        else:
            institution = '；'.join(institution_list)
            institution2 = ''

    # print(institution)

    # 6. journal
    journal_txt = filetxt.split('-----next page-----')[0:2]
    journal_txt = ''.join(journal_txt).replace(' ', '').split('\n')
    filename_journal = filepath.split('/')[-2]
    journal_txt_list = []
    journal_list = []
    for word in journal_txt:
        if filename_journal in word:
            journal_txt_list.append(word)
        if 'DOI' in word:
            journal_txt_list.append(word)
    # j name
    journal_list.append('《' + filename_journal + '》')
    # print(journal_txt)

    # j date
    qikan_list = ['No.', 'No．', 'NO.', 'DOI', '2018', '二○一八', '期', '辑', '第', '二零一八']
    for word in journal_txt_list:
        word = word.replace('，', ',').replace(' ', '').replace('．', '.').replace('·', '.')
        if 'No.' in word:
            journal_list.append('第' + word.split('No.')[1].split(',')[0] + '期')
            break
        elif 'NO.' in word:
            journal_list.append('第' + word.split('NO.')[1].split(',')[0] + '期')
            break
        elif 'DOI' in word and '2018.' in word:
            journal_list.append('第' + word.split('2018.')[1].split('.')[0] + '期')
            break
        elif re.findall('2018年第.期', word) != []:
            journal_list.append('第' + word.split('2018年第')[1].split('期')[0] + '期')
            break
        elif '2018.' in word:
            journal_list.append('第' + word.split('2018.')[1].split('.')[0] + '期')
            break

    # j year
    journal_list.append('2018年')

    # j page
    try:
        journal_page_start = filetxt.split('-----next page-----')[0].strip('\n').replace(' ', '').split()[-1]
    except Exception:
        return
    journal_page_end = filetxt.strip('\n').replace(' ', '').split()[-1]
    # print(journal_page_end)
    if journal_page_start.isdigit():
        journal_list.append('第' + journal_page_start + '页—第' + journal_page_end + '页')
    elif len(journal_page_start.split('·')) > 1:
        journal_page_start = journal_page_start.strip('·')
        journal_page_end = journal_page_end.strip('·')
        journal_list.append('第' + journal_page_start + '页—第' + journal_page_end + '页')
    elif len(journal_page_start.split('-')) > 1:
        journal_page_start = journal_page_start.strip('-')
        journal_page_end = journal_page_end.strip('-')
        journal_list.append('第' + journal_page_start + '页—第' + journal_page_end + '页')

    journal = '，'.join(journal_list)
    # print(journal)

    # 7. project
    project_txt = first_page.replace(' ','').replace('：',':')
    xiangmu_list = ['*本文为','*本文系','*本文是','本文系','本文为','本文是','本论文为','本论文系','本论文是']
    project = ''
    for word in xiangmu_list:
        project_txt = project_txt.replace(word, "*本文系")
    if re.findall("\*本文系.*?成果。", project_txt) != []:
        print(re.findall("\*本文系.*?成果。", project_txt))
        project = re.findall("\*本文系.*?成果。", project_txt)[0]
    elif '\n基金项目' in project_txt and '单..位' in project_txt:
        index1 = project_txt.find('\n基金项目')
        index2 = project_txt.find('单..位')
        if index1 < index2:
            project = project_txt.split('\n基金项目')[1].split('单..位')[0].replace('\n','')
        else :
            project_a = project_txt.split('\n基金项目')[1].split('。')[0].replace('\n','')
            if len(project_a) < 150:
                project = project_a
    elif project == '' and '\n基金项目' in project_txt:
        project = project_txt.split('\n基金项目')[1].split('\n')[0]
    elif project == '' and '基金项目' in project_txt:
        project = project_txt.split('基金项目')[1].split('\n')[0]
    else:
        project = ''
    if project != '' and project[0] == ':':
        project = project.replace(project[0],'')
    if project != '' and project[0] == '］':
        project = project.replace(project[0],'')

    project = project.replace('*','').replace(' ','')

    # 8. country
    country_max,frequency_list = country(filepath)
    country_list = str(frequency_list)

    f.close()

    outputTable.write(m,0,title)
    outputTable.write(m,1,en_title)
    outputTable.write(m,2,author)
    outputTable.write(m,3,institution)
    outputTable.write(m,4,author2)
    outputTable.write(m,5,institution2)
    outputTable.write(m,6,keywords)
    outputTable.write(m,7,abstract)
    outputTable.write(m,8,project)
    outputTable.write(m,9,journal)
    outputTable.write(m,10,country_max)
    outputTable.write(m,11,country_list)



    # # docx
    # para_title = document.add_paragraph()
    # try:
    #     run_title = para_title.add_run(title)
    # except Exception:
    #     return
    # run_title.font.name = "黑体"
    # run_title._element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")

    # para_author = document.add_paragraph()
    # run_author = para_author.add_run('【作  者】' + author)
    # run_author.font.name = '宋体'
    # run_author._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")

    # para_institution = document.add_paragraph()
    # run_institution = para_institution.add_run('【单  位】' + institution)
    # run_institution.font.name = '宋体'
    # run_institution._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")

    # para_journal = document.add_paragraph()
    # run_journal = para_journal.add_run('【期  刊】' + journal)
    # run_journal.font.name = '宋体'
    # run_journal._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")

    # para_abstract = document.add_paragraph()
    # run_abstract = para_abstract.add_run('【内容摘要】' + abstract)
    # run_abstract.font.name = '宋体'
    # run_abstract._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")

    # para_keyword = document.add_paragraph()
    # run_keyword = para_keyword.add_run('【关键词】' + keywords)
    # run_keyword.font.name = '宋体'
    # run_keyword._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")

    # document.add_page_break()

    # Excel
    # excelpath = '/Users/YiFaye_Lee/Desktop/外国文学/nianjian2018/error_find.xls'
    # Err = []
    # str = '，'

    # yesyesyes!
    # if author == '找不到【作者】！':
    #     # Err.append('找不到【作者】')
    #     errtable2.write(m,2,"找不到【作者】！")
    # if title == '找不到【标题】！':
    #     errtable2.write(m,3,"找不到【标题】！")
    # if keywords == '找不到【关键词】！':
    #     errtable2.write(m,4,"找不到【关键词】！")
    # if institution == '找不到【作者单位】！':
    #     errtable2.write(m,5,"找不到【作者单位】！")
    # if abstract == '找不到【摘要】！':
    #     errtable2.write(m,6,"找不到【摘要】！")
    # m = m + 1
    
    # if len(str.join(Err)) > 5:
    #     ErrMsg = filename + str.join(Err)
    # else:
    #     ErrMsg = ''
    # write_excel_xls_append(excelpath, ErrMsg)

# Find files that ends with .txt
def find(obj):
    if obj.endswith(".txt"):
        path_list.append(obj)


def get_list_dir(dir_path):
    dir_files = os.listdir(dir_path)
    for file in dir_files:
        file_path = os.path.join(dir_path, file)
        if os.path.isfile(file_path):
            find(file_path)
        if os.path.isdir(file_path):
            get_list_dir(file_path)




# Extract text first
pdfpath = filedialog.askdirectory(initialdir=os.getcwd(),
                                         title="请选择你想要提取的文件夹")
output(pdfpath)

# Get those txt file dir
path_list=[]
get_list_dir(pdfpath)

# read error
readErrorfile = xlwt.Workbook(encoding = 'utf-8')
errtable = readErrorfile.add_sheet('data',cell_overwrite_ok=True)
errtable.write(0,0,"ioerror title")
errtable.write(0,1,"ioerror journal")
errtable.write(0,2,"ValueError title")
errtable.write(0,3,"ValueError journal")
errtable.write(0,4,"IndexError title")
errtable.write(0,5,"IndexError journal")

# found error
# foundErrorfile = xlwt.Workbook(encoding = 'utf-8')
# errtable2 = foundErrorfile.add_sheet('data',cell_overwrite_ok=True)
# errtable.write(0,0,"论文名")
# errtable.write(0,1,"期刊名")
# errtable.write(0,2,"【作者】")
# errtable.write(0,3,"【标题】")
# errtable.write(0,4,"【关键词】")
# errtable.write(0,5,"【单位】")
# errtable.write(0,6,"【摘要】")

# Then find info
outputFile = xlwt.Workbook(encoding = 'utf-8')
outputTable = outputFile.add_sheet('data',cell_overwrite_ok=True)
outputTable.write(0,0,"中文题名")
outputTable.write(0,1,"英文题名")
outputTable.write(0,2,"第一作者（含职称职务等信息）")
outputTable.write(0,3,"机构")
outputTable.write(0,4,"第二作者（含职称职务等信息）")
outputTable.write(0,5,"机构")
outputTable.write(0,6,"关键词")
outputTable.write(0,7,"摘要")
outputTable.write(0,8,"项目")
outputTable.write(0,9,"期刊")
outputTable.write(0,10,"国别")
outputTable.write(0,11,"国别词频参考")


i = 1
j = 1
k = 1
m = 1
for path in path_list:
    try:
        filename = path.split("/")[-1]
        filename_journal = path.split('/')[-2]
        lunwen(path,m)
        m = m + 1
        print("")
    except IOError as e:
        errtable.write(i,0,filename)
        errtable.write(i,1,filename_journal)
        i = i + 1
    # except ValueError as e:
    #     errtable.write(j,2,filename)
    #     errtable.write(j,3,filename_journal)
    #     j = j + 1
    # except IndexError as e: 
    #     errtable.write(k,4,filename)
    #     errtable.write(k,5,filename_journal)
    #     k = k + 1

# document.save('/Users/YiFaye_Lee/Desktop/外国文学/nianjian2018/lunwen.docx')

outputFile.save('/Users/YiFaye_Lee/Desktop/外国文学/nianjian2018/output.xls')
readErrorfile.save('/Users/YiFaye_Lee/Desktop/外国文学/nianjian2018/error.xls')
# foundErrorfile.save('/Users/YiFaye_Lee/Desktop/外国文学/nianjian2018/error_found.xls')


